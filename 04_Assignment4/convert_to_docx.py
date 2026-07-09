#!/usr/bin/env python3
"""Convert Assignment 4 markdown to APA 7 DOCX"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
input_md = Path(r"A:\TTED731\04_Assignment4\Drafts\TTED731_Assignment4_Rafael_Rodriguez_Workforce_Assessment_Draft.md")
output_docx = Path(r"A:\TTED731\04_Assignment4\Final\TTED731_Assignment4_Workforce_Alignment_Assessment_Rafael_Rodriguez.docx")
output_docx.parent.mkdir(parents=True, exist_ok=True)

# Read markdown
with open(input_md, 'r', encoding='utf-8') as f:
    content = f.read()

# Create document
doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Process lines
lines = content.split('\n')
in_references = False

for line in lines:
    if not line.strip():
        continue
    
    if line.startswith('# '):
        # Main title
        title = line[2:].strip()
        p = doc.add_paragraph(title)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2.0
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
    
    elif line.startswith('## '):
        # Heading 2
        in_references = 'References' in line
        heading = line[3:].strip()
        p = doc.add_paragraph(heading)
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.first_line_indent = Inches(0.5)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
    
    elif line.startswith('### '):
        # Heading 3
        heading = line[4:].strip()
        p = doc.add_paragraph(heading)
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.first_line_indent = Inches(0.5)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
    
    else:
        # Body text
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        
        # References use hanging indent
        if in_references:
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
        else:
            p.paragraph_format.first_line_indent = Inches(0.5)
        
        # Process formatting
        text = line.strip()
        pos = 0
        
        for match in re.finditer(r'\*\*([^*]+)\*\*|\*([^*]+)\*', text):
            # Add text before match
            if match.start() > pos:
                run = p.add_run(text[pos:match.start()])
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            
            # Add formatted text
            if match.group(1):  # Bold
                run = p.add_run(match.group(1))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
            else:  # Italic
                run = p.add_run(match.group(2))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.italic = True
            
            pos = match.end()
        
        # Add remaining text
        if pos < len(text):
            run = p.add_run(text[pos:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

# Save
doc.save(str(output_docx))
print(f"✓ Document created: {output_docx}")
