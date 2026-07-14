import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

MD_PATH = r"a:/TTED731/05_Final_Project/Final_Project_ReadyForDOCX.md"
DOCX_PATH = r"a:/TTED731/05_Final_Project/Final_Project_ReadyForDOCX.docx"

def add_page_number(footer):
    p = footer.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._r.append(fld)


def set_doc_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)


def apply_paragraph_format(paragraph):
    pformat = paragraph.paragraph_format
    pformat.space_before = Pt(0)
    pformat.space_after = Pt(0)
    try:
        pformat.line_spacing = 2.0
    except Exception:
        pass


def parse_inline_italics(text, paragraph):
    # Split on asterisks to apply italics runs
    parts = re.split(r"(\*[^*]+\*)", text)
    for part in parts:
        if part.startswith('*') and part.endswith('*') and len(part) > 1:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def main():
    doc = Document()
    set_doc_styles(doc)

    # set 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    with open(MD_PATH, 'r', encoding='utf-8') as f:
        lines = f.read().splitlines()

    in_title_block = False
    title_lines = []
    i = 0
    # detect and extract title block up to page-break
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('<div style="text-align:center">'):
            in_title_block = True
            i += 1
            continue
        if in_title_block:
            if line.strip().startswith('</div>'):
                in_title_block = False
                # add title content centered
                for tl in title_lines:
                    if tl.strip() == '':
                        continue
                    p = doc.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # detect bold title (surrounded by **)
                    if tl.strip().startswith('**') and tl.strip().endswith('**'):
                        run = p.add_run(tl.strip().strip('*'))
                        run.bold = True
                        run.font.size = Pt(14)
                    else:
                        p.add_run(tl.strip())
                    apply_paragraph_format(p)
                # page break
                doc.add_page_break()
                title_lines = []
                i += 1
                break
            else:
                title_lines.append(line)
                i += 1
                continue
        i += 1

    # continue parsing remaining lines
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped == '':
            i += 1
            continue
        if stripped.startswith('# '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(text, style='Heading 1')
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            apply_paragraph_format(p)
            i += 1
            continue
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            p = doc.add_paragraph(text, style='Heading 2')
            apply_paragraph_format(p)
            i += 1
            continue
        if stripped.startswith('- '):
            # list item
            p = doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
            apply_paragraph_format(p)
            i += 1
            continue
        # references: detect heading 'References' and then treat following paragraphs as references
        if stripped == 'References' or stripped == '# References':
            p = doc.add_paragraph('References', style='Heading 1')
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            apply_paragraph_format(p)
            i += 1
            continue
        # Otherwise, normal paragraph; check for inline italics markers
        p = doc.add_paragraph()
        parse_inline_italics(line, p)
        apply_paragraph_format(p)
        i += 1

    # Footer page numbers
    for section in doc.sections:
        add_page_number(section.footer)

    # Apply hanging indent to references: try to find references section and apply formatting
    # Simple heuristic: find the 'References' heading paragraph index
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'References':
            # subsequent paragraphs until end are references; apply hanging indent
            for ref_para in doc.paragraphs[idx+1:]:
                if ref_para.text.strip() == '':
                    continue
                pfmt = ref_para.paragraph_format
                pfmt.left_indent = Inches(0.5)
                pfmt.first_line_indent = Inches(-0.5)
            break

    doc.save(DOCX_PATH)
    print('Saved:', DOCX_PATH)

if __name__ == '__main__':
    main()
