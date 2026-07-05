from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

md_path = Path(r"A:\TTED731\03_Assignment3_Instructional_Design\Drafts\v3\TTED731_Assignment3_Instructional_Design_Rafael_Rodriguez_Version3_Master_Revision_CitationPass.md")
out_path = Path(r"A:\TTED731\03_Assignment3_Instructional_Design\Final\TTED731_Assignment3_Instructional_Design_Rafael_Rodriguez_Final_APA7.docx")

text = md_path.read_text(encoding="utf-8")

# Remove staging HTML comment
text = re.sub(r"(?s)^<!--.*?-->\s*", "", text)

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2
style.paragraph_format.space_after = Pt(0)

# Page number in header
header = section.header
p = header.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run()
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
run._r.append(fld)

lines = text.splitlines()

# Extract title block
title = lines[0].replace("# ", "").strip()
meta = []
i = 1
while i < len(lines) and lines[i].strip() == "":
    i += 1
while i < len(lines) and not lines[i].startswith("## "):
    if lines[i].strip():
        meta.append(lines[i].strip())
    i += 1

# APA student title page
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(title)
r.bold = True

for item in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(item)

doc.add_page_break()

in_refs = False

def add_para(t, hanging=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(0)
    if hanging:
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)

    # Basic italic markdown support
    parts = re.split(r"(\*[^*]+\*)", t)
    for part in parts:
        if part.startswith("*") and part.endswith("*"):
            run = p.add_run(part.strip("*"))
            run.italic = True
        else:
            p.add_run(part)
    return p

while i < len(lines):
    line = lines[i].strip()

    if not line:
        i += 1
        continue

    if line.startswith("## References"):
        in_refs = True
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("References")
        r.bold = True
        i += 1
        continue

    if line.startswith("## "):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line.replace("## ", "").strip())
        r.bold = True
        i += 1
        continue

    if line.startswith("### "):
        p = doc.add_paragraph()
        r = p.add_run(line.replace("### ", "").strip())
        r.bold = True
        i += 1
        continue

    if in_refs:
        add_para(line, hanging=True)
    else:
        add_para(line)

    i += 1

doc.save(out_path)
print(f"Created: {out_path}")
